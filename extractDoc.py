import re
import os
import win32com.client as win32
import pandas as pd
import generateTokens
import pickle as p
import time
from glob import glob
from datetime import datetime
from docx import Document
from win32com.client import constants

def save_as_docx(folder):
	''' Converts .doc files to .docx so it can
	be read by python-docx. '''

	print('C:\\Users\\hagen\\documents\\uva\\groene-project\\' + folder.replace('/','\\') + '**\\*.doc')
	paths = glob('C:\\Users\\hagen\\documents\\uva\\groene-project\\' + folder.replace('/','\\') + '**\\*.doc', recursive=True)
	for file in paths:
		
		# Opening MS Word
		word = win32.gencache.EnsureDispatch('Word.Application')
		doc = word.Documents.Open(file)
		doc.Activate ()

		# Rename path with .docx
		new_file_abs = os.path.abspath(file)
		new_file_abs = re.sub(r'\.\w+$', '.docx', new_file_abs)

		# Save and Close
		word.ActiveDocument.SaveAs(
			new_file_abs, FileFormat=constants.wdFormatXMLDocument
		)
		doc.Close(False)
		print('Saved ' + file + ' as docx.')


def getKrantenInfo(folder):
	''' Runs through the unfortunate format of LexisNexis
	.docx files and saves the newspaper info to a csv. '''


	li_files = [file for file in os.listdir(folder) if file.endswith('.docx')]
	print(li_files)

	weekdays = ['maandag', 'dinsdag', 'woensdag', 'donderdag', 'vrijdag', 'zaterdag', 'zondag']
	months_en = ['january', 'february', 'march', 'april', 'may', 'june', 'july', 'august', 'september', 'october', 'november', 'december']
	months_nl = ['januari', 'februari', 'maart', 'april', 'mei', 'juni', 'juli', 'augustus', 'september', 'oktober', 'november', 'december']
	years = ['1995','1996','1997','1998','1999','2000','2001','2002','2003','2004','2005','2006','2007','2008','2009','2010','2011','2012','2013','2014','2015','2016','2017','2018']
	sources = ['de volkskrant', 'de telegraaf', 'ad/algemeen dagblad', 'nrc handelsblad']

	li_artikelen = []
	start_of_doc = False
	is_date = False
	is_title = False
	print_para = False
	prev_para = ''

	for index, file in enumerate(li_files):
		print(file)
		
		source = ''

		# For debugging
		# if file != 'De_Volkskrant-2006-1-moslim-moslim-islam-atleast5.docx':
		if 1 == 2:
		 	print('')
		else:
			print('Reading file ' + str(index) + ' of ' + str(len(li_files)) + ': ' +  file)
			f = open(folder + file, 'rb')
			document = Document(f)

			article_no = -1

			# Variables to get:
			# artikel_info = {}
			# artikel_info['newspaper'] = ''
			# artikel_info['title'] = ''
			# artikel_info['author'] = ''
			# artikel_info['date'] = ''
			# artikel_info['pagina'] = 0
			# artikel_info['length'] = 0
			# artikel_info['full_text'] = []

			# Loop through the paragraphs
			for i, para in enumerate(document.paragraphs):
				paragraph = para
				para_full = para.text
				para = para.text.lower().rstrip().lstrip()

				# New article!
				if not start_of_doc and 'source: ' in para:
					source = para.replace('source: ', '')

				# If a new '1 of XX DOCUMENTS' is reached, start storing article data
				if 'of' in para_full and 'DOCUMENTS' in para_full:

					#print(article_no)
					if article_no >= 0:
						print('Finished article on paragraph ' + str(i) + '/' + str(len(document.paragraphs)) + ', ' + li_artikelen[article_no]['newspaper'] + ', ' + li_artikelen[article_no]['date_formatted'] + ': ' + li_artikelen[article_no]['title'])

					li_artikelen.append({})
					article_no = article_no + 1
					li_artikelen[article_no]['newspaper'] = source
					li_artikelen[article_no]['title'] = ''
					li_artikelen[article_no]['full_text'] = []
					start_of_doc = True
					is_title = False
					is_date = False

				# Only store data after the start of a doc
				if start_of_doc:

					if len(para) > 0:

						# Only catch the first newspaper header. Sometimes newspaper names appear as in-text headers.
						if para in sources and ('of' in prev_para and 'DOCUMENTS' in prev_para):
							is_date = True
							li_artikelen[article_no]['newspaper'] = para

						elif is_date:
							if 'load-date' in para or '1vetkolom' in para:
								is_date = True
							else:
								#print('Published on ' + para_full)
								li_artikelen[article_no]['date'] = para
								date_formatted = para.replace(',','').rstrip()
								for year in years:
									if year in date_formatted:
										date_formatted = date_formatted.split(year)
										date_formatted = date_formatted[0] + year
								for z, day in enumerate(weekdays):
									if day in date_formatted:
										date_formatted = date_formatted.replace(' ' + day, '')
								for x, month in enumerate(months_nl):
									if month in date_formatted:
										date_formatted = date_formatted.replace(month, months_en[x])
								if date_formatted[0].isdigit():
									li_artikelen[article_no]['date_formatted'] = datetime.strptime(date_formatted.rstrip(), '%d %B %Y').strftime('%Y-%m-%d')
								else:
									li_artikelen[article_no]['date_formatted'] = datetime.strptime(date_formatted.rstrip(), '%B %d %Y').strftime('%Y-%m-%d')

							is_date = False
							# Title (always?) comes after date
							is_title = True

						# Each of these comes after the title. If these are encountered, the title section is over.
						elif 'SECTION: ' in para_full:
							li_artikelen[article_no]['pagina'] = para.lower().replace('section: ', '')
							is_title = False
						elif 'LENGTH: ' in para_full:
							li_artikelen[article_no]['length'] = int(para.lower().replace('length: ', '').replace(' words', '').replace(' woorden', ''))
							is_title = False
						elif 'BYLINE: ' in para_full:
							li_artikelen[article_no]['author'] = para.lower().replace('byline: ', '')
							is_title = False
						elif 'SOURCE: ' in para_full:
							is_title = False

						elif is_title:
							
							if para_full in weekdays:
								is_title = True
							else:
								#print('New article on paragraph ' + str(i) + '/' + str(len(document.paragraphs)) + ': "' + para_full + '"')
								if li_artikelen[article_no]['title'] != '':
									li_artikelen[article_no]['title'] = ' '.join([li_artikelen[article_no]['title'], para_full])
								else:
									li_artikelen[article_no]['title'] = para_full
								#is_title = False
							# A subtitle is incoming when the title ends with '; ', so set is_title to True again	

						elif is_title == False and len(para) > 0 and para != ' ' and not para_full.endswith(' DOCUMENTS') and 'Time Of Request' not in para_full and 'All Documents:' not in para_full and 'LANGUAGE:' not in para_full and 'PUBLICATION-TYPE:' not in para_full and 'All Rights Reserved' not in para_full and not para_full.startswith('SECTION: ') and not para_full.startswith('GRAPHIC: ') and not para_full.startswith('JOURNAL-CODE: ') and not para_full.startswith('LOAD-DATE: ') and not para_full.startswith('Copyright 2') and not is_title:

							li_artikelen[article_no]['full_text'].append(para_full)
							# print('Regular text paragraph')
							# print(para_full)

						prev_para = para_full

					# End of doc
					if i == (len(document.paragraphs) - 1):
						print('End of doc')
						print('Writing results to csv')

						print(type(li_artikelen[article_no]['full_text']),li_artikelen[article_no]['full_text'])
						# Join the list of paragraphs with a newline (can be used as paragraph split later)
						li_artikelen[article_no]['full_text'] = '\n'.join(li_artikelen[article_no]['full_text'])
						print(type(li_artikelen[article_no]['full_text']),li_artikelen[article_no]['full_text'])

						start_of_doc = True

						# Write to csv
						columns = ['date','date_formatted','author','newspaper','title','full_text','pagina','length']
						df = pd.DataFrame(li_artikelen, columns=columns)
						print(df.head())
						# Write to file
						if os.path.isfile(folder + 'all-data.csv'):
							df.to_csv(folder + 'all-data.csv', mode='a', header=False)
						else:
							df.to_csv(folder + 'all-data.csv')

						li_artikelen = []
						start_of_doc = False
						is_date = False
						is_title = False
						article_no = -1
						f.close()

			#print(li_artikelen)
			#print(document)
			

#save_as_docx('shit')
if __name__ == '__main__':
	
	#getKrantenInfo('data/media/kranten/moslim-islam/')
	# tokens = getTokens.getNewspaperTokens('data/media/kranten/multicultureel-multiculturele-multiculturalisme/all-data.csv')
	# p.dump(tokens, open('data/media/kranten/tokens-all-multicultureel-multiculturele-multiculturalisme.p', 'wb'))